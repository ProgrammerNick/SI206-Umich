# Import packages
# You need to install the package scikit-learn in order for lines referring to sklearn to run
# You need to install the package python-docx in order for lines referring to docx to run
import pandas as pd
import statsmodels.api as sm
import numpy as np
from sklearn.linear_model import Lasso
import datetime
import matplotlib.pyplot as plt
import sklearn.metrics as metrics
from sklearn.tree import DecisionTreeRegressor
from sklearn.tree import plot_tree
from sklearn import tree
from sklearn.tree import _tree
from docx import Document

# Change the number of rows and columns to display
pd.set_option('display.max_rows', 200)
pd.set_option('display.min_rows', 200)
pd.set_option('display.max_columns', 50)
pd.set_option('display.max_colwidth', 50)
pd.set_option('display.precision', 5)
pd.options.display.float_format = '{:.5f}'.format

# Define a path for import and export
path = 'C:/Users/Nicolas Newberry/OneDrive/Documents/Umich Assignments/FIN427/'

# Import data
returns01 = pd.read_csv(path + 'Aggregate data 20250120_1144.csv')
returns01['month'] = pd.to_datetime(returns01['month'], format='%d-%b-%Y')
print(returns01.dtypes)
print(returns01.head(10))
print(returns01.columns)

# Decision-tree analysis
y = returns01['indadjret']
x = returns01[['lag1mcreal','finlag1bm','finlag1bmmiss']]
dt = DecisionTreeRegressor(max_depth=3, min_weight_fraction_leaf=0.10)
fn = ['lag1mcreal','finlag1bm','finlag1bmmiss']
dtmodel = dt.fit(x, y)
dtresult = dt.score(x, y)
dtpredictions = dt.predict(x)
dfscore = pd.DataFrame([dtresult])
dfpredictions = pd.DataFrame([dtpredictions])
dfreview = pd.DataFrame({'Name': returns01['comnam'],
                         'cusip9': returns01['cusip9'],
                         'month': returns01['month'],
                         'lag1mcreal': returns01['lag1mcreal'],
                         'finlag1bm': returns01['finlag1bm'],
                         'finlag1bmmiss': returns01['finlag1bmmiss'],
                         'indadjret': y,
                         'pred_indadjret': dtpredictions})
print(dfreview)

plot_tree(dtmodel, feature_names=fn)
plt.savefig(path + 'Basic decision tree 20250129_1231.pdf')
plt.show()

# What is the relative importance of each feature in lowering mean squared error?
# There is tabular output and a chart, the chart does not display well because of the number of features.
# I have left the chart code in the file in case you want it for another project.
importances = dt.feature_importances_
sorted_index = np.argsort(importances)[::-1]
ximportance = range(len(importances))
labels = np.array(fn)[sorted_index]
plt.bar(ximportance, importances[sorted_index], tick_label=labels)
plt.xticks(rotation=90)
plt.savefig(path + 'Basic decision tree importance 20250129_1232.pdf')
plt.show()

dfimportance = pd.DataFrame(list(zip(labels, importances[sorted_index])), columns=['fn', 'importances'])

print('Decision-tree score:', dtresult)
print('Mean Absolute Error:', metrics.mean_absolute_error(y, dtpredictions))
print('Mean Squared Error:', metrics.mean_squared_error(y, dtpredictions))
print('Root Mean Squared Error:', np.sqrt(metrics.mean_squared_error(y, dtpredictions)))
print('Features sorted by importance:', dfimportance)

# Write rules to text or Word documents
feature_names = list(x.columns)

# Write rules to a text document
def get_rules(tree, feature_names, spacer_base="    ", output_file=None):
    left = tree.tree_.children_left
    right = tree.tree_.children_right
    threshold = tree.tree_.threshold
    features = [
        feature_names[i] if i != _tree.TREE_UNDEFINED else "undefined!"
        for i in tree.tree_.feature
    ]
    value = tree.tree_.value

    lines = []

    def recurse(node, depth):
        spacer = spacer_base * depth
        if threshold[node] != _tree.TREE_UNDEFINED:
            line = f"{spacer}if {features[node]} <= {threshold[node]:.4f}:"
            lines.append(line)
            recurse(left[node], depth + 1)
            line = f"{spacer}else:  # if {features[node]} > {threshold[node]:.4f}"
            lines.append(line)
            recurse(right[node], depth + 1)
        else:
            line = f"{spacer}return {value[node]}"
            lines.append(line)

    recurse(0, 0)

    if output_file:
        with open(output_file, 'w') as f:
            for line in lines:
                f.write(line + '\n')

# Export rules to a text file
get_rules(dtmodel, feature_names, output_file=path + 'Example 04 Decision tree rules 20250129_1232.txt')

# Function to write rules to a Word document
def get_rules_word(tree, feature_names, document, spacer_base='    '):
    left = tree.tree_.children_left
    right = tree.tree_.children_right
    threshold = tree.tree_.threshold
    features = [
        feature_names[i] if i != _tree.TREE_UNDEFINED else 'undefined!'
        for i in tree.tree_.feature
    ]
    value = tree.tree_.value

    def recurse(node, depth):
        spacer = spacer_base * depth
        if threshold[node] != _tree.TREE_UNDEFINED:
            document.add_paragraph(f'{spacer}if {features[node]} <= {threshold[node]:.5f}:')
            recurse(left[node], depth + 1)
            document.add_paragraph(f'{spacer}else:  # if {features[node]} > {threshold[node]:.5f}')
            recurse(right[node], depth + 1)
        else:
            document.add_paragraph(f'{spacer}return {value[node]}')

    recurse(0, 0)

# Export rules to a Word document
doc = Document()
get_rules_word(dtmodel, feature_names, doc)
doc.save(path + 'Example 04 Decision tree rules 20250129_1232.docx')

# Create a dataframe of the last month to illustrate predicted returns
examplemonth = datetime.datetime(2023, 12, 31)
dfexample = dfreview[dfreview['month'] == examplemonth]
print(dfexample)
print(dfexample.dtypes)

# Export results to Excel
with pd.ExcelWriter(path + 'Excel 04 Decision-tree output 20250129_1232.xlsx') as writer:
    dfscore.to_excel(writer, sheet_name='Score')
    dfimportance.to_excel(writer, sheet_name='Importance')
    dfexample.to_excel(writer, sheet_name='Example')