# Import packages
# You need to install the package scikit-learn in order for lines referring to sklearn to run
# You need to install the package python-docx in order for lines referring to docx to run
import pandas as pd
import statsmodels.api as sm
import numpy as np
from sklearn.linear_model import Lasso
import datetime
import matplotlib.pyplot as plt
import sklearn.metrics as metrics
from sklearn.tree import DecisionTreeRegressor
from sklearn.tree import plot_tree
from sklearn import tree
from sklearn.tree import _tree
from docx import Document

# Change the number of rows and columns to display
pd.set_option('display.max_rows', 200)
pd.set_option('display.min_rows', 200)
pd.set_option('display.max_columns', 50)
pd.set_option('display.max_colwidth', 50)
pd.set_option('display.precision', 5)
pd.options.display.float_format = '{:.5f}'.format

# Define a path for import and export
path = 'C:/Users/Nicolas Newberry/OneDrive/Documents/Umich Assignments/fin427/'

# Import and view data
returns01 = pd.read_csv(path + 'Aggregate data 20250208_1500.csv')
returns01['month'] = pd.to_datetime(returns01['month'], format='%d-%b-%Y')
print(returns01.dtypes)
print(returns01.head(10))
print(returns01.columns)

# Compile descriptive statistics
# For continuous variables show the mean, standard deviation and percentiles
# For indicator (1,0) variables show the mean
mean        =  returns01[['ret', 'adjret', 'gsret', 'indadjret', 'lag1mcreal', 'adjlag1bm', 'g06_invpegadj']].mean()
stddev      =  returns01[['ret', 'adjret', 'gsret', 'indadjret', 'lag1mcreal', 'adjlag1bm', 'g06_invpegadj']].std()
count       =  returns01[['ret', 'adjret', 'gsret', 'indadjret', 'lag1mcreal', 'adjlag1bm', 'g06_invpegadj']].count()
percentiles = (returns01[['ret', 'adjret', 'gsret', 'indadjret', 'lag1mcreal', 'adjlag1bm', 'g06_invpegadj']]
               .quantile([0, 0.125, 0.500, 0.875, 1]))
meanind     = returns01[['finlag1bmmiss','fing06_invpegadjmiss']].mean()
print(mean)
print(stddev)
print(count)
print(percentiles)
print(meanind)

# What are the correlations between variables
correlation_matrix01 = returns01[['indadjret',  'lnlag1mcreal',
                                   'adjlag1bm'   ,  'finlag1bm'       , 'finlag1bmmiss',
                                  'g06_invpegadj',  'fing06_invpegadj', 'fing06_invpegadjmiss']].corr()
correlation_matrix02 = returns01[['indadjret', 'zlnlag1mcreal',
                                  'zadjlag1bm'    ,'zfinlag1bm'       , 'finlag1bmmiss',
                                  'zg06_invpegadj','zfing06_invpegadj', 'fing06_invpegadjmiss']].corr()
correlation_matrix03 = returns01[['indadjret', 'rlnlag1mcreal',
                                  'radjlag1bm'    , 'rfinlag1bm'      , 'finlag1bmmiss',
                                  'rg06_invpegadj','rfing06_invpegadj', 'fing06_invpegadjmiss']].corr()
print(correlation_matrix01)
print(correlation_matrix02)
print(correlation_matrix03)

# Regression using statsmodels
y = returns01['indadjret']
x = returns01[['lnlag1mcreal','finlag1bm','finlag1bmmiss','fing06_invpegadj','fing06_invpegadjmiss']]

x = sm.add_constant(x)
model = sm.OLS(y, x).fit()
predictions = model.predict(x)
print_model = model.summary()
b_coef = model.params
b_err = model.bse

# Calculating predicted values and Cook's D influence statistic, and merging with original dataset
influence = model.get_influence()
cooks_d = influence.cooks_distance[0]
analysis = pd.DataFrame({'predictions': predictions, 'cooks_d': cooks_d})
analysis = analysis.sort_values(by='cooks_d', ascending=False)

print(print_model)
print(f'R-squared: {model.rsquared:.5f}')
print(b_coef)
print(b_err)

# Computations to understand what R-squared means
y_pred = model.predict(x)
ssr = np.sum((y - y_pred)**2)
sst = np.sum((y - np.mean(y))**2)
rsq = 1 - (ssr/sst)
rmse = np.sqrt(np.mean((y - y_pred)**2))

print(f'Sum of squared difference between y values and predicted y values (SSR): {ssr:.5f}')
print(f'Sum of squared difference between y values and average y values (SST): {sst:.5f}')
print(f'R-squared = 1 - SSR/SST: {rsq:.5f}')
print(f'Square root of the mean squared error: {rmse:.5f}')

# Export to Excel
with pd.ExcelWriter(path + 'Excel 05 Descriptive and OLS 20250208_1525.xlsx') as writer:
    mean.to_excel(writer, sheet_name='mn')
    stddev.to_excel(writer, sheet_name='sd')
    count.to_excel(writer, sheet_name='n')
    percentiles.to_excel(writer, sheet_name='perc')
    meanind.to_excel(writer, sheet_name='mnind')
    correlation_matrix01.to_excel(writer, sheet_name='corr01')
    correlation_matrix02.to_excel(writer, sheet_name='corr02')
    correlation_matrix03.to_excel(writer, sheet_name='corr03')
    b_coef.to_excel(writer, sheet_name='b_coef')
    b_err.to_excel(writer, sheet_name='b_err')

# OLS regression using z-scores and LASSO regression

# Regression using statsmodels
yz = returns01['indadjret']
xz = returns01[['zlnlag1mcreal','zfinlag1bm','finlag1bmmiss','zfing06_invpegadj','fing06_invpegadjmiss']]

xz = sm.add_constant(xz)
modelz = sm.OLS(yz, xz).fit()
predictionsz = modelz.predict(xz)
print_modelz = modelz.summary()
b_coefz = modelz.params
b_errz = modelz.bse

# Calculating predicted values and Cook's D influence statistic, and merging with original dataset
influencez = model.get_influence()
cooks_dz = influencez.cooks_distance[0]
analysisz = pd.DataFrame({'predictions': predictionsz, 'cooks_d': cooks_dz})
analysisz = analysisz.sort_values(by='cooks_d', ascending=False)

print(print_modelz)
print(f'R-squared: {modelz.rsquared:.5f}')
print(b_coefz)
print(b_errz)

# Computations to understand what R-squared means
y_predz = modelz.predict(xz)
ssrz = np.sum((yz - y_predz)**2)
sstz = np.sum((yz - np.mean(yz))**2)
rsqz = 1 - (ssrz/sstz)
rmsez = np.sqrt(np.mean((yz - y_predz)**2))

print(f'Sum of squared difference between y values and predicted y values (SSR): {ssrz:.5f}')
print(f'Sum of squared difference between y values and average y values (SST): {sstz:.5f}')
print(f'R-squared = 1 - SSR/SST: {rsqz:.5f}')
print(f'Square root of the mean squared error: {rmsez:.5f}')

# LASSO regression

# Specify the penalty factor
lasso = Lasso(alpha=0.0004)

# LASSO regression
yl = returns01['indadjret']
xl = returns01[['zlnlag1mcreal','zfinlag1bm','finlag1bmmiss','zfing06_invpegadj','fing06_invpegadjmiss']]
lasso.fit(xl, yl)
lasso_coef = lasso.fit(xl, yl).coef_
lasso_score = lasso.score(xl, yl)
print(lasso.intercept_)
print(lasso_coef)
print(xl.columns)
print(pd.Series(lasso_coef, index=xl.columns))
y_pred_lasso = lasso.predict(xl)
ssr_lasso = np.sum((yl - y_pred_lasso)**2)
sstl      = np.sum((yl - np.mean(yl))**2)
rsq_lasso = 1 - (ssr_lasso/sstl)
rmse_lasso = np.sqrt(np.mean((yl - y_pred_lasso)**2))
print(f'Sum of squared difference between y values and predicted y values (SSR): {ssr_lasso:.5f}')
print(f'Sum of squared difference between y values and average y values (SST): {sstl:.5f}')
print(f'R-squared = 1 - SSR/SST: {rsq_lasso:.5f}')
print(f'Square root of the mean squared error: {rmse_lasso:.5f}')

# Create a dataframe of coefficients
lasso_coef_df = pd.DataFrame({'Feature': ['intercept'] + xl.columns.tolist(),
                        'Coefficient': [lasso.intercept_] + lasso.coef_.tolist()})
print(lasso_coef_df)

# Export to Excel
with pd.ExcelWriter(path + 'Excel 05 OLS and LASSO 20250208_1626.xlsx') as writer:
    b_coef.to_excel(writer, sheet_name='b_coef')
    b_err.to_excel(writer, sheet_name='b_err')
    lasso_coef_df.to_excel(writer, sheet_name='coef_lasso')

# Decision-tree analysis
yd = returns01['indadjret']
xd = returns01[['lag1mcreal','finlag1bm','finlag1bmmiss','fing06_invpegadj','fing06_invpegadjmiss']]
dt = DecisionTreeRegressor(max_depth=3, min_weight_fraction_leaf=0.08)
fn = ['lag1mcreal','finlag1bm','finlag1bmmiss','fing06_invpegadj','fing06_invpegadjmiss']
dtmodel = dt.fit(xd, yd)
dtresult = dt.score(xd, yd)
dtpredictions = dt.predict(xd)
dfscore = pd.DataFrame([dtresult])
dfpredictions = pd.DataFrame([dtpredictions])
dfreviewdt = pd.DataFrame({'Name': returns01['comnam'],
                         'cusip9': returns01['cusip9'],
                         'month': returns01['month'],
                         'lag1mcreal': returns01['lag1mcreal'],
                         'finlag1bm': returns01['finlag1bm'],
                         'finlag1bmmiss': returns01['finlag1bmmiss'],
                         'fing06_invpegadj': returns01['fing06_invpegadj'],
                         'ffing06_invpegadjmiss': returns01['fing06_invpegadjmiss'],
                         'indadjret': yd,
                         'pred_indadjret': dtpredictions})

plot_tree(dtmodel, feature_names=fn)
plt.savefig(path + 'Basic decision tree 20250208_1635.pdf')
plt.show()

# What is the relative importance of each feature in lowering mean squared error?
# There is tabular output and a chart, the chart does not display well because of the number of features.
# I have left the chart code in the file in case you want it for another project.
importances = dt.feature_importances_
sorted_index = np.argsort(importances)[::-1]
ximportance = range(len(importances))
labels = np.array(fn)[sorted_index]
plt.bar(ximportance, importances[sorted_index], tick_label=labels)
plt.xticks(rotation=90)
plt.savefig(path + 'Basic decision tree importance 20250208_1641.pdf')
plt.show()

dfimportance = pd.DataFrame(list(zip(labels, importances[sorted_index])), columns=['fn', 'importances'])

print('Decision-tree score:', dtresult)
print('Mean Absolute Error:', metrics.mean_absolute_error(y, dtpredictions))
print('Mean Squared Error:', metrics.mean_squared_error(y, dtpredictions))
print('Root Mean Squared Error:', np.sqrt(metrics.mean_squared_error(y, dtpredictions)))
print('Features sorted by importance:', dfimportance)

# Write rules to text or Word documents
feature_names = list(xd.columns)

# Write rules to a text document
def get_rules(tree, feature_names, spacer_base="    ", output_file=None):
    left = tree.tree_.children_left
    right = tree.tree_.children_right
    threshold = tree.tree_.threshold
    features = [
        feature_names[i] if i != _tree.TREE_UNDEFINED else "undefined!"
        for i in tree.tree_.feature
    ]
    value = tree.tree_.value

    lines = []

    def recurse(node, depth):
        spacer = spacer_base * depth
        if threshold[node] != _tree.TREE_UNDEFINED:
            line = f"{spacer}if {features[node]} <= {threshold[node]:.4f}:"
            lines.append(line)
            recurse(left[node], depth + 1)
            line = f"{spacer}else:  # if {features[node]} > {threshold[node]:.4f}"
            lines.append(line)
            recurse(right[node], depth + 1)
        else:
            line = f"{spacer}return {value[node]}"
            lines.append(line)

    recurse(0, 0)

    if output_file:
        with open(output_file, 'w') as f:
            for line in lines:
                f.write(line + '\n')

# Export rules to a text file
get_rules(dtmodel, feature_names, output_file=path + 'Example 05 Decision tree rules 20250208_1642.txt')

# Function to write rules to a Word document
def get_rules_word(tree, feature_names, document, spacer_base='    '):
    left = tree.tree_.children_left
    right = tree.tree_.children_right
    threshold = tree.tree_.threshold
    features = [
        feature_names[i] if i != _tree.TREE_UNDEFINED else 'undefined!'
        for i in tree.tree_.feature
    ]
    value = tree.tree_.value

    def recurse(node, depth):
        spacer = spacer_base * depth
        if threshold[node] != _tree.TREE_UNDEFINED:
            document.add_paragraph(f'{spacer}if {features[node]} <= {threshold[node]:.5f}:')
            recurse(left[node], depth + 1)
            document.add_paragraph(f'{spacer}else:  # if {features[node]} > {threshold[node]:.5f}')
            recurse(right[node], depth + 1)
        else:
            document.add_paragraph(f'{spacer}return {value[node]}')

    recurse(0, 0)

# Export rules to a Word document
doc = Document()
get_rules_word(dtmodel, feature_names, doc)
doc.save(path + 'Example 05 Decision tree rules 20250208_1644.docx')

# Create a dataframe of the last month to illustrate predicted returns
examplemonth = datetime.datetime(2023, 12, 31)
dfexample = dfreviewdt[dfreviewdt['month'] == examplemonth]
print(dfexample.dtypes)

# Export results to Excel
with pd.ExcelWriter(path + 'Excel 05 Decision-tree output 20250208_1645.xlsx') as writer:
    dfscore.to_excel(writer, sheet_name='Score')
    dfimportance.to_excel(writer, sheet_name='Importance')
    dfexample.to_excel(writer, sheet_name='Example')
