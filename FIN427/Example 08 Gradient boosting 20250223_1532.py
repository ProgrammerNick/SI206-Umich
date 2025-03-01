# Import packages
# You need to install the package scikit-learn in order for lines referring to sklearn to run
# You need to install the package python-docx in order for lines referring to docx to run
import pandas as pd
import statsmodels.api as sm
import numpy as np
from sklearn.linear_model import Lasso
import datetime
import matplotlib.pyplot as plt
import sklearn.metrics as metrics
from sklearn.tree import DecisionTreeRegressor
from sklearn.tree import plot_tree
from sklearn import tree
from sklearn.tree import _tree
from docx import Document
from sklearn.ensemble import RandomForestRegressor
from sklearn.ensemble import GradientBoostingRegressor
from sklearn.model_selection import ParameterGrid

# Change the number of rows and columns to display
pd.set_option('display.max_rows', 200)
pd.set_option('display.min_rows', 200)
pd.set_option('display.max_columns', 50)
pd.set_option('display.max_colwidth', 50)
pd.set_option('display.precision', 5)
pd.options.display.float_format = '{:.5f}'.format

# Define a path for import and export
path = 'E:\\UM\Courses\\FIN427\\FIN427 Winter 2025\\'

# Import and view data
returns01 = pd.read_csv(path + 'Aggregate data 20250223_1217.csv')
returns01['month'] = pd.to_datetime(returns01['month'], format='%d-%b-%Y')
print(returns01.dtypes)
print(returns01.head(10))
print(returns01.columns)

datecut = datetime.datetime(2005, 12, 31)
print(datecut)
returns01_train = returns01[returns01['month'] <= datecut]
returns01_valid = returns01[returns01['month'] >  datecut]

# Decision-tree analysis
yd = returns01['indadjret']
xd = returns01[['lag1mcreal','finlag1bm','finlag1bmmiss',
                'fing02esg','fing02esgmiss',
                'fing04fcfyadj','fing04fcfyadjmiss',
                'fing06_invpegadj','fing06_invpegadjmiss',
                'fing11ret5adj','fing11ret5adjmiss',
                'fing13sueadj','fing13sueadjmiss',
                'fing14erevadj','fing14erevadjmiss']]

dt = DecisionTreeRegressor(max_depth=5, min_weight_fraction_leaf=0.08)
fn = ['lag1mcreal','finlag1bm','finlag1bmmiss',
                'fing02esg','fing02esgmiss',
                'fing04fcfyadj','fing04fcfyadjmiss',
                'fing06_invpegadj','fing06_invpegadjmiss',
                'fing11ret5adj','fing11ret5adjmiss',
                'fing13sueadj','fing13sueadjmiss',
                'fing14erevadj','fing14erevadjmiss']
dtmodel = dt.fit(xd, yd)
dtresult = dt.score(xd, yd)
dtpredictions = dt.predict(xd)
dfdtscore = pd.DataFrame([dtresult])
dfpredictions = pd.DataFrame([dtpredictions])
dfreviewdt = pd.DataFrame({'Name': returns01['comnam'],
                         'cusip9': returns01['cusip9'],
                         'month': returns01['month'],
                         'lag1mcreal': returns01['lag1mcreal'],
                         'finlag1bm': returns01['finlag1bm'],
                         'finlag1bmmiss': returns01['finlag1bmmiss'],
                         'fing02esg': returns01['fing02esg'],
                         'fing02esgmiss': returns01['fing02esgmiss'],
                         'fing04fcfyadj': returns01['fing04fcfyadj'],
                         'fing04fcfyadjmiss': returns01['fing04fcfyadjmiss'],
                         'fing06_invpegadj': returns01['fing06_invpegadj'],
                         'fing06_invpegadjmiss': returns01['fing06_invpegadjmiss'],
                         'fing11ret5adj': returns01['fing11ret5adj'],
                         'fing11ret5adjmiss': returns01['fing11ret5adjmiss'],
                         'fing13sueadj': returns01['fing13sueadj'],
                         'fing13sueadjmiss': returns01['fing13sueadjmiss'],
                         'fing14erevadj': returns01['fing14erevadj'],
                         'fing14erevadjmiss': returns01['fing14erevadjmiss'],
                         'indadjret': yd,
                         'pred_indadjret': dtpredictions})
plot_tree(dtmodel, feature_names=fn)
plt.savefig(path + 'Basic decision tree 20250224_1221.pdf')
plt.show()

# What is the relative importance of each feature in lowering mean squared error?
# There is tabular output and a chart, the chart does not display well because of the number of features.
# I have left the chart code in the file in case you want it for another project.
dtimportances = dt.feature_importances_
dtsorted_index = np.argsort(dtimportances)[::-1]
xdtimportance = range(len(dtimportances))
dtlabels = np.array(fn)[dtsorted_index]
plt.bar(xdtimportance, dtimportances[dtsorted_index], tick_label=dtlabels)
plt.xticks(rotation=90)
plt.savefig(path + 'Basic decision tree importance 20250224_1221.pdf')
plt.show()

dfdtimportance = pd.DataFrame(list(zip(dtlabels, dtimportances[dtsorted_index])), columns=['fn', 'dtimportances'])

print('Decision-tree score:', dtresult)
print('Mean Absolute Error:', metrics.mean_absolute_error(yd, dtpredictions))
print('Mean Squared Error:', metrics.mean_squared_error(yd, dtpredictions))
print('Root Mean Squared Error:', np.sqrt(metrics.mean_squared_error(yd, dtpredictions)))
print('Features sorted by importance:', dfdtimportance)

# Write rules to text or Word documents
feature_names = list(xd.columns)

# Write rules to a text document
def get_rules(tree, feature_names, spacer_base="    ", output_file=None):
    left = tree.tree_.children_left
    right = tree.tree_.children_right
    threshold = tree.tree_.threshold
    features = [
        feature_names[i] if i != _tree.TREE_UNDEFINED else "undefined!"
        for i in tree.tree_.feature
    ]
    value = tree.tree_.value

    lines = []

    def recurse(node, depth):
        spacer = spacer_base * depth
        if threshold[node] != _tree.TREE_UNDEFINED:
            line = f"{spacer}if {features[node]} <= {threshold[node]:.4f}:"
            lines.append(line)
            recurse(left[node], depth + 1)
            line = f"{spacer}else:  # if {features[node]} > {threshold[node]:.4f}"
            lines.append(line)
            recurse(right[node], depth + 1)
        else:
            line = f"{spacer}return {value[node]}"
            lines.append(line)

    recurse(0, 0)

    if output_file:
        with open(output_file, 'w') as f:
            for line in lines:
                f.write(line + '\n')

# Export rules to a text file
get_rules(dtmodel, feature_names, output_file=path + 'Example 07 Decision tree rules 20250224_1221.txt')

# Function to write rules to a Word document
def get_rules_word(tree, feature_names, document, spacer_base='    '):
    left = tree.tree_.children_left
    right = tree.tree_.children_right
    threshold = tree.tree_.threshold
    features = [
        feature_names[i] if i != _tree.TREE_UNDEFINED else 'undefined!'
        for i in tree.tree_.feature
    ]
    value = tree.tree_.value

    def recurse(node, depth):
        spacer = spacer_base * depth
        if threshold[node] != _tree.TREE_UNDEFINED:
            document.add_paragraph(f'{spacer}if {features[node]} <= {threshold[node]:.5f}:')
            recurse(left[node], depth + 1)
            document.add_paragraph(f'{spacer}else:  # if {features[node]} > {threshold[node]:.5f}')
            recurse(right[node], depth + 1)
        else:
            document.add_paragraph(f'{spacer}return {value[node]}')

    recurse(0, 0)

# Export rules to a Word document
doc = Document()
get_rules_word(dtmodel, feature_names, doc)
doc.save(path + 'Example 07 Decision tree rules 20250224_1221.docx')

# Create a dataframe of the last month to illustrate predicted returns
examplemonth = datetime.datetime(2023, 12, 31)
dfdtexample = dfreviewdt[dfreviewdt['month'] == examplemonth]
print(dfdtexample.dtypes)

# Random forests
# Above, we split on ALL possible features, and specified the maximum depth and minimum leaf percentage.
# Let's generate many trees in which we specify the maximum number of features under consideration
# and continue to specify several different maximum depth and minimum leaf percentage.
# When we specify the maximum number of features, in each forest a random selection of features will be considered
# in each tree. This will be useful when there is a very large number of features for possible consideration.
# In our analysis we started with a small enough number of features and observations such that the computer could
# execute the analysis in reasonable time. But if we had thousands of features and millions of observations,
# consideration of all the features at once could be problematic.

# First, run a random forest with a specified maximum depth, minimum weight fraction in each leaf, number of trees
# and maximum features. There is a random state assumption as well, which allows us to check results. If
# we run the program multiple times with the same random state we should get the same result.
# Later, we will consider multiple potential hyperparameters at once.

# We are going to compile aggregate results as well as result for each tree generated, to illustrate
# the process. We will keep the score from each decision tree as well as feature importances

# Prepare an empty dataframes to store feature importances
importancesx = pd.DataFrame()
# Prepare an empty list of scores
scores_list = []
# The array was needed to generate feature names for output to dataframes.
xd_array = xd.values

# Specify hyperparameters
rf = RandomForestRegressor(max_depth=5, min_weight_fraction_leaf=0.08, n_estimators=100, random_state=11610,
                           max_features=8, bootstrap=True, max_samples=0.60)
rfmodel = rf.fit(xd_array, yd)
rfresult_full = rf.score(xd_array, yd)
rfscore_full = pd.DataFrame([rfresult_full])
rfimportances_full = rf.feature_importances_
rfsorted_index_full = np.argsort(rfimportances_full)[::-1]
rflabels_full = np.array(fn)[rfsorted_index_full]
rfimportance_full = pd.DataFrame(list(zip(rflabels_full, rfimportances_full[rfsorted_index_full])), columns=['fn', 'importances'])
rfpredictions_full = rf.predict(xd_array)
rfreview_full = pd.DataFrame({'month': returns01['month'],'cusip9': returns01['cusip9'],
'comnam': returns01['comnam'],
'indadjret': yd,
'pred_indadjret': rfpredictions_full})

# Iterate over each Decision Tree and collect feature importances
for i, tree in enumerate(rf.estimators_):
    scorex = tree.score(xd_array, yd)
    scores_list.append({'Tree': i+1, 'Score': scorex})

    temp_df = pd.DataFrame({
        'Feature': xd.columns,
        'Importance': tree.feature_importances_,
        'Tree': i+1})
    importancesx = pd.concat([importancesx, temp_df], ignore_index=True)

scores_df = pd.DataFrame(scores_list)
scores_df.reset_index(drop=True, inplace=True)
importancesx.reset_index(drop=True, inplace=True)

print('Decision-tree score:', rfresult_full)
print('Features sorted by importance:', rfimportance_full)
print(scores_df)
# print(importancesx)
dfrfexample = rfreview_full[rfreview_full['month'] == examplemonth]

# What if we could run many trees, but each iteration of the tree used information from the prior tree to lower
# estimation error. This is gradient boosting.
# Number of trees: Search over large number of (short) trees
# Maximum depth: The idea is that a large number of short trees may perform well.
# Learning rate: Each subsequent tree counts less in the predicted output.
# Subsample: Portion of sample on which to estimate each tree (1 = 100%; 0.5 = 50% etc).

# Prepare an empty dataframes to store feature importances
importancesxgb = pd.DataFrame()
# Prepare an empty list of scores
scores_listgb = []
# The array was needed to generate feature names for output to dataframes.
xd_arraygb = xd.values

gb = GradientBoostingRegressor(max_depth=5, min_weight_fraction_leaf=0.08, n_estimators=100, random_state=11610,
                               max_features=8, subsample=0.60,  learning_rate=0.1)
gbmodel_full = gb.fit(xd_arraygb, yd)
gbresult_full = gb.score(xd_arraygb, yd)
gbscore_full = pd.DataFrame([gbresult_full])
gbimportances_full = gb.feature_importances_
gbsorted_index_full = np.argsort(gbimportances_full)[::-1]
gblabels_full = np.array(fn)[gbsorted_index_full]
gbimportance_full = pd.DataFrame(list(zip(gblabels_full, gbimportances_full[gbsorted_index_full])), columns=['fn', 'importances'])
gbpredictions_full = gb.predict(xd_arraygb)

gbreview_full = pd.DataFrame({'month': returns01['month'],'cusip9': returns01['cusip9'],
'comnam': returns01['comnam'],
'indadjret': yd,
'pred_indadjret': gbpredictions_full})

# Iterate over each Decision Tree and collect feature importances
for i, tree in enumerate(gb.estimators_):
    decision_tree = tree[0]
    scorexgb = decision_tree.score(xd_arraygb, yd)
    scores_listgb.append({'Tree': i+1, 'Score': scorexgb})

    temp_dfgb = pd.DataFrame({
        'Feature': xd.columns,
        'Importance': decision_tree.feature_importances_,
        'Tree': i+1})
    importancesxgb = pd.concat([importancesxgb, temp_dfgb], ignore_index=True)

scores_dfgb = pd.DataFrame(scores_listgb)
scores_dfgb.reset_index(drop=True, inplace=True)
importancesxgb.reset_index(drop=True, inplace=True)

print('Gradient boosted score:', gbresult_full)
print('Features sorted by importance:', gbimportance_full)
print(scores_dfgb)
# print(importancesxgb)
dfgbexample = gbreview_full[gbreview_full['month'] == examplemonth]

# Export results to Excel
with pd.ExcelWriter(path + 'Excel 08 Decision-tree output 20250224_1402.xlsx') as writer:
    dfdtscore.to_excel(writer, sheet_name='DT Score')
    dfdtimportance.to_excel(writer, sheet_name='DT Importance')
    dfdtexample.to_excel(writer, sheet_name='DT Example')
    rfscore_full.to_excel(writer, sheet_name='Score full rf')
    rfimportance_full.to_excel(writer, sheet_name='Importance full rf')
    scores_df.to_excel(writer, sheet_name='Scores by iteration rf')
    importancesx.to_excel(writer, sheet_name='Importance by iteration rf')
    dfrfexample.to_excel(writer, sheet_name='RF Example')
    gbscore_full.to_excel(writer, sheet_name='Score full gb')
    gbimportance_full.to_excel(writer, sheet_name='Importance full gb')
    scores_dfgb.to_excel(writer, sheet_name='Scores by iteration gb')
    importancesxgb.to_excel(writer, sheet_name='Importance by iteration gb')
    dfgbexample.to_excel(writer, sheet_name='GB Example')

